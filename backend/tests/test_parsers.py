"""Tests for document parsers."""
import io
import textwrap

import pytest

from app.parsers.docx_parser import TextParser
from app.parsers import SUPPORTED_MIME_TYPES


# ── TextParser ────────────────────────────────────────────────────────────────

class TestTextParser:
    def setup_method(self):
        self.parser = TextParser()

    def test_plain_text_single_page(self):
        content = b"Hello world"
        doc = self.parser.parse(content)
        assert len(doc.pages) == 1
        assert doc.pages[0].text == "Hello world"
        assert doc.pages[0].page_number == 1

    def test_long_text_splits_into_pages(self):
        # TextParser splits on 3000 chars
        content = ("A" * 3001).encode()
        doc = self.parser.parse(content)
        assert len(doc.pages) >= 2

    def test_empty_content(self):
        doc = self.parser.parse(b"")
        assert len(doc.pages) == 1
        assert doc.pages[0].text == ""

    def test_utf8_decoding(self):
        text = "Ångström – 日本語テスト"
        doc = self.parser.parse(text.encode("utf-8"))
        assert text in doc.pages[0].text

    def test_mime_type_text_plain(self):
        assert "text/plain" in SUPPORTED_MIME_TYPES

    def test_mime_type_markdown(self):
        assert "text/markdown" in SUPPORTED_MIME_TYPES


# ── SUPPORTED_MIME_TYPES ──────────────────────────────────────────────────────

class TestSupportedMimeTypes:
    def test_pdf_supported(self):
        assert "application/pdf" in SUPPORTED_MIME_TYPES

    def test_docx_supported(self):
        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in SUPPORTED_MIME_TYPES

    def test_image_types_supported(self):
        for mime in ("image/jpeg", "image/png", "image/webp"):
            assert mime in SUPPORTED_MIME_TYPES

    def test_unsupported_raises(self):
        from app.parsers import parse_document
        import asyncio
        with pytest.raises(ValueError, match="Unsupported MIME type"):
            asyncio.get_event_loop().run_until_complete(
                parse_document(b"data", "application/octet-stream")
            )


# ── TableAwarePDFParser — smoke test with minimal PDF bytes ───────────────────

class TestTableAwarePDFParser:
    def test_invalid_bytes_raises(self):
        from app.parsers.table_parser import TableAwarePDFParser
        parser = TableAwarePDFParser()
        with pytest.raises(Exception):
            parser.parse(b"not a pdf")

    def test_parse_returns_parsed_document(self):
        """Create a minimal valid 1-page PDF in memory with PyMuPDF."""
        import fitz
        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from test PDF")
        doc.save(buf)
        buf.seek(0)

        from app.parsers.table_parser import TableAwarePDFParser
        parser = TableAwarePDFParser()
        result = parser.parse(buf.read())

        assert len(result.pages) == 1
        assert "Hello from test PDF" in result.pages[0].text


# ── DOCXParser ────────────────────────────────────────────────────────────────

class TestDOCXParser:
    def test_invalid_bytes_raises(self):
        from app.parsers.docx_parser import DOCXParser
        parser = DOCXParser()
        with pytest.raises(Exception):
            parser.parse(b"not a docx")

    def test_parse_real_docx(self):
        """Build a minimal DOCX with python-docx and parse it."""
        from docx import Document as DocxDocument
        from app.parsers.docx_parser import DOCXParser

        buf = io.BytesIO()
        docx = DocxDocument()
        docx.add_heading("Introduction", level=1)
        docx.add_paragraph("This is paragraph one.")
        docx.add_paragraph("This is paragraph two.")
        docx.save(buf)
        buf.seek(0)

        parser = DOCXParser()
        result = parser.parse(buf.read())

        full_text = " ".join(p.text for p in result.pages)
        assert "Introduction" in full_text
        assert "paragraph one" in full_text
