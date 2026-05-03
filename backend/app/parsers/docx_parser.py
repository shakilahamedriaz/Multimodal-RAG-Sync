import io

from docx import Document as DocxDocument

from app.schemas import ParsedDocument, ParsedPage

_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"}
_LOGICAL_PAGE_CHARS = 3000  # ~750 tokens, enough context per simulated page


class DOCXParser:
    def parse(self, content: bytes) -> ParsedDocument:
        doc = DocxDocument(io.BytesIO(content))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name in _HEADING_STYLES:
                lines.append(f"\n## {text}")
            else:
                lines.append(text)

        full_text = "\n".join(lines).strip()

        # DOCX has no real page breaks — split into logical pages by char count
        pages: list[ParsedPage] = []
        for i in range(0, max(1, len(full_text)), _LOGICAL_PAGE_CHARS):
            chunk = full_text[i : i + _LOGICAL_PAGE_CHARS].strip()
            if chunk:
                pages.append(ParsedPage(page_number=len(pages) + 1, text=chunk))

        if not pages:
            pages = [ParsedPage(page_number=1, text="")]

        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return ParsedDocument(pages=pages, page_count=len(pages), mime_type=mime)


class TextParser:
    def parse(self, content: bytes) -> ParsedDocument:
        text = content.decode("utf-8", errors="replace").strip()
        return ParsedDocument(
            pages=[ParsedPage(page_number=1, text=text)],
            page_count=1,
            mime_type="text/plain",
        )
